# -*- coding: utf-8 -*-
"""
generate_manual.py
產生 RAS400 公差分析 AI 系統說明書（Word + PDF）
"""

import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, PageBreak, HRFlowable)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

OUTPUT_DIR = r"C:\Tolerance_Project\說明書"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── 字型設定（使用系統內建標楷體）──────────────────────────────────────────
FONT_PATHS = [
    r"C:\Windows\Fonts\kaiu.ttf",       # 標楷體
    r"C:\Windows\Fonts\msjh.ttc",       # 微軟正黑體
    r"C:\Windows\Fonts\mingliu.ttc",    # 細明體
]
FONT_NAME = None
for fp in FONT_PATHS:
    if os.path.exists(fp):
        try:
            pdfmetrics.registerFont(TTFont("CJK", fp))
            FONT_NAME = "CJK"
            break
        except Exception:
            continue
if not FONT_NAME:
    # fallback: Helvetica (no CJK)
    FONT_NAME = "Helvetica"
    print("[WARN] 找不到中文字型，PDF 可能無法顯示中文")

# ════════════════════════════════════════════════════════════════
# 說明書內容定義
# ════════════════════════════════════════════════════════════════

TITLE   = "RAS400 公差分析 AI 系統說明書"
SUBTITLE = "Web-based Tolerance Analysis System with AI Assistant"
VERSION  = "v2.0   |   2026-05"

SECTIONS = [
    # ── 1 系統概述 ──────────────────────────────────────────────
    {
        "heading": "一、系統概述",
        "body": [
            "本系統為「RAS400 精密迴轉滑台公差分析 AI 系統」，以 Web 介面取代傳統 Excel/tkinter 單機版，透過瀏覽器即可執行公差累積分析、敏感度／貢獻度計算、蒙地卡羅模擬、機台媒合與 AI 智能對話等功能。",
            "系統採用 Flask 後端（Python）、Ollama 本地 LLM（gemma3:4b）或雲端 API 作為對話引擎，並整合 Neo4j 知識圖譜（RAS400 本體庫）提供精準的零件配合建議。",
        ],
        "list": [
            "啟動方式：執行 run_ai.bat，瀏覽器開啟 http://127.0.0.1:7011",
            "預設模型：gemma3:4b（本地 Ollama）",
            "分析核心：Jacobian 矩陣法 + 統計公差（RSS）+ 蒙地卡羅模擬",
            "組裝目標：RAS400 C 軸精密迴轉滑台（11 個零件）",
        ],
    },
    # ── 2 介面說明 ───────────────────────────────────────────────
    {
        "heading": "二、介面說明",
        "body": [
            "系統介面由三個區域組成：左側功能按鈕面板、中央圖形顯示區（同步即時更新）、右側 AI 對話框。",
        ],
        "subsections": [
            {
                "heading": "2.1 左側功能按鈕面板",
                "table": {
                    "headers": ["按鈕名稱", "功能說明", "觸發方式"],
                    "rows": [
                        ["產品架構圖", "顯示 RAS400 11 個零件的 BOM 樹狀結構圖", "點擊 → AI 回覆 + 彈窗"],
                        ["產品架構圖（特徵）", "顯示各零件的特徵面（P/S/H/C）架構圖", "點擊 → AI 回覆 + 彈窗"],
                        ["公差網路圖", "顯示零件間的公差累積網路拓樸圖", "點擊 → AI 回覆 + 彈窗"],
                        ["組裝接觸圖", "顯示零件接觸配合關係圖（綠色連線）", "點擊 → AI 回覆 + 彈窗"],
                        ["STEP 3D 檢視器", "開啟 STEP 檔案 3D 瀏覽面板（需上傳 STP）", "直接開啟側邊面板"],
                        ["編輯公差路徑", "開啟公差路徑編輯器，可新增/修改路徑項目", "直接開啟 Modal"],
                        ["公差分析", "開啟公差分析參數設定並執行分析", "送 AI 訊息處理"],
                        ["公差調配", "開啟調配參數 Modal（自動調配 / 手動比對）", "直接開啟 Modal"],
                        ["製程與機台媒合", "開啟機台搜尋與配合建議介面", "直接開啟 Modal"],
                        ["配合建議", "開啟 RAS400 零件配合建議比對面板", "直接開啟 Modal"],
                        ["進階版配合建議", "開啟配合建議報告產生 + 路徑更新 Wizard", "直接開啟面板"],
                    ]
                }
            },
            {
                "heading": "2.2 右側對話框",
                "body": [
                    "右側為 AI 對話區，使用者可輸入自然語言問題或指令，系統結合 RAS400 知識圖譜與公差分析結果給予回覆。",
                    "上方可切換 AI 模型（本地 Ollama / 雲端 API），右上角顯示 Model 選單。",
                ]
            }
        ]
    },
    # ── 3 AI 對話指令對照表 ───────────────────────────────────────
    {
        "heading": "三、AI 對話指令對照表",
        "body": [
            "以下列出輸入關鍵字後系統的觸發行為與預期輸出，分為「直接觸發」與「AI 對話」兩大類。",
        ],
        "subsections": [
            {
                "heading": "3.1 圖表類指令",
                "table": {
                    "headers": ["輸入範例", "觸發類型", "預期輸出"],
                    "rows": [
                        ["RAS400 有哪些零件", "AI + BOM 圖", "零件清單文字 + 左側 BOM 樹狀圖彈窗"],
                        ["畫出產品架構圖", "AI + BOM 圖", "11 個零件的層次架構圖"],
                        ["畫出公差網路圖", "AI + 網路圖", "零件-特徵-公差連線拓樸圖彈窗"],
                        ["畫出組裝接觸關係圖", "AI + 接觸圖", "零件間接觸配合綠色連線圖"],
                        ["畫出所有零件的特徵面結構圖", "AI + 特徵圖", "各零件 P/S/H/C 特徵架構圖彈窗"],
                        ["顯示目前的公差路徑", "AI 純文字", "列出目前路徑項目的代號與數值"],
                    ]
                }
            },
            {
                "heading": "3.2 公差分析類指令",
                "table": {
                    "headers": ["輸入範例", "觸發類型", "預期輸出"],
                    "rows": [
                        ["執行公差分析", "開啟分析 Modal", "彈出蒙地卡羅參數設定視窗"],
                        ["進行公差分析", "開啟分析 Modal", "同上"],
                        ["公差累積分析結果", "AI 純文字", "解讀目前的 RSS / WC 誤差數值"],
                        ["公差調配", "開啟調配 Modal", "彈出自動調配（四象限）參數設定"],
                        ["公差最佳化", "開啟調配 Modal", "同上"],
                    ]
                }
            },
            {
                "heading": "3.3 配合建議類指令",
                "table": {
                    "headers": ["輸入範例", "觸發類型", "預期輸出"],
                    "rows": [
                        ["軸承配合", "Fast Path 配合建議", "RAS400 軸承的 YRT 外圈 H6 / 內圈 js5 配合建議"],
                        ["軸承座配合", "Fast Path 配合建議", "軸承座孔公差 H6 / H7 建議選項"],
                        ["工作臺心軸配合", "Fast Path 配合建議", "心軸軸頸配合建議（H7/k6 等）"],
                        ["需要定位精確可裝拆的配合", "ANSI 維度搜尋", "符合「定位＋可裝拆」的標準配合代號"],
                        ["H7/h6 是什麼配合", "AI 純文字", "間隙配合說明、公差帶計算"],
                        ["25mm H7/h6 的間隙是多少", "AI 純文字", "依 ISO 286 計算間隙範圍"],
                        ["配合建議", "開啟配合建議 Modal", "直接開啟 RAS400 配合比對面板"],
                        ["進階版", "開啟進階版配合建議面板", "配合建議報告 + 路徑更新 Wizard"],
                    ]
                }
            },
            {
                "heading": "3.4 製程與機台類指令",
                "table": {
                    "headers": ["輸入範例", "觸發類型", "預期輸出"],
                    "rows": [
                        ["搜尋適合的機台", "開啟媒合 Modal", "彈出製程與機台媒合視窗"],
                        ["製程", "AI 製程建議", "針對公差等級推薦加工製程（磨削/鏜孔等）"],
                        ["加工方式", "AI 製程建議", "同上"],
                        ["磨削、車削、銑削...", "AI 製程建議", "說明對應加工方法的公差能力"],
                    ]
                }
            },
            {
                "heading": "3.5 路徑編輯類指令",
                "table": {
                    "headers": ["輸入範例", "觸發類型", "預期輸出"],
                    "rows": [
                        ["編輯公差路徑", "開啟編輯器", "彈出路徑編輯器 Modal"],
                        ["修改路徑", "開啟編輯器", "同上"],
                        ["安插路徑項目", "開啟編輯器 + AI 說明", "說明如何安插新的路徑元素"],
                        ["放寬 IT 等級", "開啟編輯器 + AI 調整", "AI 建議放寬哪個公差項目"],
                        ["收緊 IT7 → IT6", "開啟編輯器 + AI 調整", "AI 執行公差等級調整"],
                        ["選用軸承配合 H6/js5", "軸承配合套用", "自動更新對應路徑項目的公差值"],
                    ]
                }
            },
            {
                "heading": "3.6 注意事項（不應觸發分析的指令）",
                "table": {
                    "headers": ["輸入範例", "系統行為", "說明"],
                    "rows": [
                        ["25mm H7/h6 配合分析", "AI 純文字（不開分析 Modal）", "含「配合」→ 視為 ISO 查詢，非路徑分析"],
                        ["軸承配合分析", "AI 純文字", "含「分析」且含「配合」→ 排除分析 Modal"],
                        ["查看路徑", "AI 純文字（不開編輯器）", "含「查看」→ 視為只讀意圖"],
                        ["公差調整", "AI 純文字", "須用「公差調配」完整詞組才開 Modal"],
                    ]
                }
            }
        ]
    },
    # ── 4 公差分析操作流程 ───────────────────────────────────────
    {
        "heading": "四、公差分析操作流程",
        "body": [
            "以下為完整的公差分析操作步驟，建議依序執行以獲得最佳分析結果。",
        ],
        "steps": [
            ("Step 1：匯入公差路徑",
             "點擊「編輯公差路徑」按鈕開啟路徑編輯器。\n"
             "可透過「匯入 Excel/CSV」載入已有路徑，或手動新增路徑項目（平移 traX/Y/Z、旋轉 rotX/Y/Z 與各類幾何公差）。\n"
             "路徑代碼格式：零件名稱-公差類型-編號（例：軸承座-Dia-1）。"),
            ("Step 2：設定分析參數",
             "點擊「公差分析」按鈕或在對話框輸入「執行公差分析」。\n"
             "設定蒙地卡羅取樣數（預設 10,000）、標準差倍數（2σ / 3σ / 4σ）、分布模式（均勻 / 常態）。\n"
             "點擊「執行公差分析」等待計算完成。"),
            ("Step 3：查看分析結果",
             "分析完成後，左側切換「分析報告」面板，包含：\n"
             "・統計摘要：RSS ±3σ、Worst Case、MC 標準差（X/Y/Z mm，aX/aY/aZ arc_sec）\n"
             "・誤差分佈：6 個軸向的蒙地卡羅直方圖\n"
             "・敏感度：各公差項對誤差方向的敏感度百分比（Q1-Q4 四象限顏色）\n"
             "・貢獻度：各公差項的實際誤差貢獻百分比\n"
             "・3D 散點圖：6DOF 誤差空間分佈（可旋轉互動）"),
            ("Step 4：公差調配（選用）",
             "點擊「公差調配」按鈕，選擇調配模式：\n"
             "・自動調配：選定目標軸向、目標 RSS 值與策略（精度優先/成本優先），系統自動計算最佳分配。\n"
             "・手動比對：在路徑編輯器手動修改公差值後，比對與基準分析的改善幅度（%）。\n"
             "調配完成後，報告顯示 Q1-Q4 四象限診斷與各軸改善百分比。"),
            ("Step 5：匯出報表",
             "點擊「導出報表」按鈕可下載 Excel 格式的完整分析報告，包含 Tideal 矩陣、RSS 結果、敏感度排名、貢獻度排名等資料。"),
        ]
    },
    # ── 5 分析結果說明 ───────────────────────────────────────────
    {
        "heading": "五、分析結果說明",
        "subsections": [
            {
                "heading": "5.1 誤差輸出項目",
                "table": {
                    "headers": ["項目", "說明", "單位"],
                    "rows": [
                        ["RSS ±Nσ", "均方根誤差（N 倍標準差範圍），統計公差的主要指標", "mm（位移）/ arc_sec（角度）"],
                        ["Worst Case (WC)", "最惡情況誤差（所有公差取最大值相加），代表絕對上限", "mm / arc_sec"],
                        ["MC Std (σ)", "蒙地卡羅模擬的標準差（N 次取樣統計）", "mm / arc_sec"],
                        ["MC Max", "蒙地卡羅模擬中出現的最大絕對誤差", "mm / arc_sec"],
                        ["Tideal Matrix", "理想（無公差）狀態的齊次轉換矩陣（4×4）", "—"],
                    ]
                }
            },
            {
                "heading": "5.2 四象限診斷指標",
                "body": [
                    "公差項目依敏感度（%）與貢獻度（%）的中位數分為四個象限，顯示於敏感度／貢獻度圖表與調配報告中：",
                ],
                "table": {
                    "headers": ["象限", "特性", "建議動作", "顏色"],
                    "rows": [
                        ["Q1", "高敏感 × 高貢獻", "優先收緊（對改善誤差效果最大）", "紅色"],
                        ["Q2", "高敏感 × 低貢獻", "規格嚴守，不可放寬", "藍色"],
                        ["Q3", "低敏感 × 高貢獻", "公差過大，收緊即可大幅改善", "橘色"],
                        ["Q4", "低敏感 × 低貢獻", "可考慮放寬以降低加工成本", "綠色"],
                    ]
                }
            }
        ]
    },
    # ── 6 配合建議功能 ───────────────────────────────────────────
    {
        "heading": "六、配合建議功能說明",
        "subsections": [
            {
                "heading": "6.1 配合建議 Modal（plan_compare）",
                "body": [
                    "點擊「配合建議」按鈕，開啟 RAS400 11 個零件的標準配合比對面板。",
                    "使用者選擇零件後，系統顯示該零件的配合選項（孔公差代號 / 軸公差代號）、功能需求標籤、配合類型（間隙 / 過渡 / 過盈）。",
                    "選擇配合後點擊「套用」，系統自動更新路徑編輯器中對應的公差值。",
                ],
                "table": {
                    "headers": ["零件", "配合代號", "配對對象", "用途"],
                    "rows": [
                        ["工作臺(1)", "H7/k6", "工作臺心軸(5)", "定位精確可裝拆"],
                        ["軸承座(2)", "H6", "軸承YRT(3) 外圈", "固定定位"],
                        ["軸承座(2)", "H7/h6", "馬達水套(7)", "定位可裝拆"],
                        ["軸承YRT(3)", "H6 / js5", "軸承座(2) / 工作臺心軸(5)", "外圈固定 / 內圈過渡"],
                        ["轉動軸(4)", "H7/u6", "工作臺心軸(5)", "強力壓入永久固定"],
                        ["工作臺心軸(5)", "js5", "軸承(3)內圈", "精密過渡配合"],
                        ["馬達(6)", "H7/s6 / H7/h6", "馬達水套(7) / 馬達座(10)", "中壓入 / 定位可裝拆"],
                    ]
                }
            },
            {
                "heading": "6.2 進階版配合建議",
                "body": [
                    "點擊「進階版配合建議」按鈕，進入兩模式介面：",
                    "【配合建議報告】：當公差分析已完成時，自動讀取分析結果，依 Q1-Q4 象限生成各公差項的配合優先建議（排序：Q1 > Q3 > Q2 > Q4）。",
                    "【更新公差路徑 Wizard】：Step 1 確認路徑狀態 → Step 2 選擇基準零件 → Step 3 選擇配合代號 → Step 4 完成，直接進入公差調配。",
                ]
            }
        ]
    },
    # ── 7 製程與機台媒合 ─────────────────────────────────────────
    {
        "heading": "七、製程與機台媒合功能",
        "body": [
            "點擊「製程與機台媒合」按鈕，開啟媒合介面，分為兩個步驟：",
        ],
        "subsections": [
            {
                "heading": "Step 1：Smart Fit 配合搜尋",
                "body": [
                    "輸入功能需求關鍵字（如「定位精確」「可裝拆」「強制壓入」），並輸入加工直徑（mm）。",
                    "設定製程能力指標 Cp（1.0 基本 / 1.33 一般 / 1.67 嚴格 / 3.0 高精密），系統依 IT 等級計算所需機台精度。",
                    "點擊「搜尋適合的配合 & 機台」，系統從 ANSI B4.1 資料庫回傳匹配的配合代號。",
                ],
            },
            {
                "heading": "Step 2：製程與機台篩選",
                "body": [
                    "設定應用場景（通用 / 模具 / 航太 / 量產）、公司、機台類型（VMC / HMC / 五軸 / 車床 / 磨床等）。",
                    "輸入目標重現性與定位精度（可由 Step 1 自動帶入），或設定行程、轉速等進階篩選條件。",
                    "點擊「開始篩選機台」，系統從機台資料庫篩選符合條件的機台，依綜合推薦 Score、規格符合度、資料信賴度排序。",
                    "點擊「產生機台媒合報表」可同步至 AI 記憶，後續對話可直接詢問媒合結果。",
                ],
            }
        ]
    },
    # ── 8 路徑代碼對照表 ─────────────────────────────────────────
    {
        "heading": "八、公差路徑代碼對照表",
        "body": [
            "路徑代碼由「零件名稱-公差類型-編號」組成，公差類型代碼如下：",
        ],
        "table": {
            "headers": ["代碼前綴", "公差類型", "方向", "範例"],
            "rows": [
                ["traX/Y/Z", "理想平移（名義尺寸）", "X / Y / Z 方向", "traY（沿 Y 軸平移 149mm）"],
                ["rotX/Y/Z", "理想旋轉（名義角度）", "繞 X / Y / Z 軸", "rotZ（繞 Z 軸旋轉 90°）"],
                ["disX/Y/Z", "距離公差（線性）", "X / Y / Z 方向", "軸承座-Dis-1（disZ）"],
                ["flaX/Y/Z", "平面度公差", "X / Y / Z 方向", "工作臺心軸-Fla-1（flaZ）"],
                ["symX/Y/Z", "對稱度公差", "X / Y / Z 方向", "—"],
                ["AngX/Y/Z", "傾斜度公差（單軸）", "繞 X / Y / Z 軸", "—"],
                ["PerX/Y/Z", "垂直度公差（單軸）", "繞 X / Y / Z 軸", "工作臺心軸-Per-1（PerX）"],
                ["ParX/Y/Z", "平行度公差（單軸）", "繞 X / Y / Z 軸", "軸承-Par-1（ParX）"],
                ["CraX/Y/Z", "軸向圓偏轉（單軸）", "繞 X / Y / Z 軸", "—"],
                ["dia", "直徑公差", "徑向", "軸承座-Dia-2（dia）"],
                ["rad", "半徑公差", "徑向", "—"],
                ["cir", "真圓度公差", "徑向", "軸承座-Cir-1（cir）"],
                ["cy", "圓柱度公差", "徑向", "—"],
                ["pos", "位置度公差", "徑向", "—"],
                ["str", "真直度公差", "徑向", "—"],
                ["con / co", "同心度公差", "徑向", "—"],
                ["crd", "徑向圓偏轉度", "徑向", "—"],
            ]
        }
    },
    # ── 9 注意事項 ───────────────────────────────────────────────
    {
        "heading": "九、注意事項與常見問題",
        "subsections": [
            {
                "heading": "9.1 系統啟動",
                "list": [
                    "每次執行 run_ai.bat 前，建議先執行「taskkill /F /IM python.exe」確保無殘留舊程序",
                    "Ollama 需先啟動並載入模型（gemma3:4b 為預設免費模型）",
                    "AI 首次回覆可能需要 10~30 秒載入知識圖譜",
                ],
            },
            {
                "heading": "9.2 模型選擇",
                "list": [
                    "llama3:8b-instruct-q4_K_M 現已需要 Ollama 付費訂閱，系統已自動從選單移除",
                    "免費本地模型：gemma3:4b（推薦）、llama3.1:8b",
                    "若回覆品質不佳，可嘗試切換至雲端模型（需設定 API Key）",
                ],
            },
            {
                "heading": "9.3 分析常見問題",
                "list": [
                    "公差分析完成後，角度誤差（aX/aY/aZ）單位為「arc_second（角秒）」，非弧度",
                    "蒙地卡羅取樣數越大，結果越精確，但計算時間越長（10,000 約 5~10 秒）",
                    "公差路徑中「角度類公差」（Par/Per/Ang/Cra）需填入正確的「轉換距離」（D 欄），用於量綱換算",
                ],
            }
        ]
    }
]


# ════════════════════════════════════════════════════════════════
# Word 文件生成
# ════════════════════════════════════════════════════════════════

def set_cell_background(cell, color_hex):
    """設定表格儲存格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_table_to_doc(doc, table_data, header_color='1F497D'):
    """加入格式化表格"""
    headers = table_data['headers']
    rows    = table_data['rows']
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_background(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        run.font.name = '微軟正黑體'

    # data rows
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(9)
            run.font.name = '微軟正黑體'
            if r_idx % 2 == 1:
                set_cell_background(cell, 'F2F2F2')
    return t

def generate_word():
    doc = Document()

    # 頁面設定
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    # 樣式設定
    normal_style = doc.styles['Normal']
    normal_style.font.name = '微軟正黑體'
    normal_style.font.size = Pt(10)

    # ── 封面 ──
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = '微軟正黑體'

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(SUBTITLE)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    run.font.name = '微軟正黑體'

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run(VERSION)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run.font.name = '微軟正黑體'

    doc.add_paragraph()
    doc.add_paragraph()

    # ── 各節 ──
    for sec in SECTIONS:
        # 一級標題
        h1 = doc.add_paragraph(sec['heading'])
        h1.style = 'Heading 1'
        run = h1.runs[0]
        run.font.name = '微軟正黑體'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # 本文
        for body_line in sec.get('body', []):
            p = doc.add_paragraph(body_line)
            p.runs[0].font.name = '微軟正黑體'
            p.runs[0].font.size = Pt(10)

        # 列表
        for item in sec.get('list', []):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(item)
            run.font.name = '微軟正黑體'
            run.font.size = Pt(10)

        # 步驟
        for step_title, step_body in sec.get('steps', []):
            h = doc.add_paragraph()
            run = h.add_run(step_title)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            run.font.name = '微軟正黑體'
            for line in step_body.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.paragraph_format.left_indent = Cm(0.7)
                    p.runs[0].font.name = '微軟正黑體'
                    p.runs[0].font.size = Pt(10)

        # 直屬表格
        if 'table' in sec:
            add_table_to_doc(doc, sec['table'])
            doc.add_paragraph()

        # 子節
        for sub in sec.get('subsections', []):
            h2 = doc.add_paragraph(sub['heading'])
            h2.style = 'Heading 2'
            run = h2.runs[0]
            run.font.name = '微軟正黑體'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

            for body_line in sub.get('body', []):
                p = doc.add_paragraph(body_line)
                p.runs[0].font.name = '微軟正黑體'
                p.runs[0].font.size = Pt(10)

            for item in sub.get('list', []):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item)
                run.font.name = '微軟正黑體'
                run.font.size = Pt(10)

            if 'table' in sub:
                add_table_to_doc(doc, sub['table'])
                doc.add_paragraph()

    out_path = os.path.join(OUTPUT_DIR, "RAS400_公差分析AI系統說明書.docx")
    doc.save(out_path)
    print(f"[OK] Word 說明書已儲存：{out_path}")
    return out_path


# ════════════════════════════════════════════════════════════════
# PDF 生成
# ════════════════════════════════════════════════════════════════

def build_pdf_styles():
    styles = getSampleStyleSheet()
    F = FONT_NAME

    base = ParagraphStyle('base', fontName=F, fontSize=10, leading=16,
                          textColor=colors.HexColor('#222222'))
    title_style = ParagraphStyle('title', parent=base, fontSize=20, leading=28,
                                 alignment=1, textColor=colors.HexColor('#1F497D'),
                                 spaceAfter=6)
    subtitle_style = ParagraphStyle('subtitle', parent=base, fontSize=13, alignment=1,
                                    textColor=colors.HexColor('#404040'), spaceAfter=4)
    h1 = ParagraphStyle('h1', parent=base, fontSize=14, leading=20, spaceBefore=14,
                         spaceAfter=6, textColor=colors.HexColor('#1F497D'),
                         borderPad=4)
    h2 = ParagraphStyle('h2', parent=base, fontSize=12, leading=18, spaceBefore=10,
                         spaceAfter=4, textColor=colors.HexColor('#2E74B5'))
    h3 = ParagraphStyle('h3', parent=base, fontSize=11, leading=16, spaceBefore=8,
                         spaceAfter=3, textColor=colors.HexColor('#2E74B5'))
    body = ParagraphStyle('body', parent=base, fontSize=10, leading=16,
                           spaceAfter=4, leftIndent=0)
    bullet = ParagraphStyle('bullet', parent=base, fontSize=10, leading=15,
                              leftIndent=14, spaceAfter=2,
                              bulletIndent=4, bulletFontName=F)
    step_title = ParagraphStyle('steptitle', parent=base, fontSize=11, leading=16,
                                 spaceBefore=8, spaceAfter=2,
                                 textColor=colors.HexColor('#1F497D'),
                                 fontName=F)
    caption = ParagraphStyle('caption', parent=base, fontSize=9, textColor=colors.grey)
    return dict(title=title_style, subtitle=subtitle_style, h1=h1, h2=h2, h3=h3,
                body=body, bullet=bullet, step_title=step_title, caption=caption, base=base)


def build_pdf_table(table_data, styles):
    F = FONT_NAME
    headers = table_data['headers']
    rows    = table_data['rows']
    col_w   = [A4[0] - 4*cm] * 1  # total width placeholder

    data = [[Paragraph(f'<b>{h}</b>', ParagraphStyle('th', fontName=F, fontSize=9,
                        textColor=colors.white, alignment=1))
             for h in headers]]
    for row in rows:
        data.append([Paragraph(str(c), ParagraphStyle('td', fontName=F, fontSize=9,
                                leading=13)) for c in row])

    n_cols = len(headers)
    avail  = A4[0] - 4.5*cm
    col_widths = [avail / n_cols] * n_cols

    tbl = Table(data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F497D')),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('FONTNAME',   (0,0), (-1,-1), F),
        ('FONTSIZE',   (0,0), (-1,-1), 9),
        ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F2F2F2')]),
        ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
    ]))
    return tbl


def generate_pdf():
    out_path = os.path.join(OUTPUT_DIR, "RAS400_公差分析AI系統說明書.pdf")
    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=2.5*cm, rightMargin=2.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    S = build_pdf_styles()
    story = []
    F = FONT_NAME

    # ── 封面 ──
    story += [
        Spacer(1, 1.5*cm),
        Paragraph(TITLE, S['title']),
        Paragraph(SUBTITLE, S['subtitle']),
        Paragraph(VERSION, ParagraphStyle('ver', fontName=F, fontSize=11, alignment=1,
                                           textColor=colors.grey)),
        Spacer(1, 0.5*cm),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1F497D')),
        Spacer(1, 0.5*cm),
    ]

    # ── 各節 ──
    for sec in SECTIONS:
        story.append(Paragraph(sec['heading'], S['h1']))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                 color=colors.HexColor('#1F497D'), spaceAfter=4))

        for body_line in sec.get('body', []):
            story.append(Paragraph(body_line, S['body']))

        for item in sec.get('list', []):
            story.append(Paragraph(f'• {item}', S['bullet']))

        for step_title, step_body in sec.get('steps', []):
            story.append(Paragraph(step_title, S['step_title']))
            for line in step_body.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, ParagraphStyle('step_body', fontName=F,
                                            fontSize=10, leading=15, leftIndent=14)))

        if 'table' in sec:
            story.append(Spacer(1, 0.2*cm))
            story.append(build_pdf_table(sec['table'], S))
            story.append(Spacer(1, 0.3*cm))

        for sub in sec.get('subsections', []):
            story.append(Paragraph(sub['heading'], S['h2']))

            for body_line in sub.get('body', []):
                story.append(Paragraph(body_line, S['body']))

            for item in sub.get('list', []):
                story.append(Paragraph(f'• {item}', S['bullet']))

            if 'table' in sub:
                story.append(Spacer(1, 0.2*cm))
                story.append(build_pdf_table(sub['table'], S))
                story.append(Spacer(1, 0.3*cm))

        story.append(Spacer(1, 0.2*cm))

    doc.build(story)
    print(f"[OK] PDF 說明書已儲存：{out_path}")
    return out_path


# ════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print("=== 產生 RAS400 公差分析 AI 系統說明書 ===")
    generate_word()
    generate_pdf()
    print("=== 完成！輸出目錄：", OUTPUT_DIR, "===")
